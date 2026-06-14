from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BODY_FONT = "宋体"
HEADING_FONT = "黑体"
CODE_FONT = "Consolas"
MATH_FONT = "Cambria Math"

BODY_SIZE = Pt(12)   # 小四
HEADING_SIZE = Pt(15)  # 小三
TITLE_SIZE = Pt(18)
SUBTITLE_SIZE = Pt(15)
CODE_SIZE = Pt(10.5)

FLOWCHART_IMAGE_MAP = {
    "图3-1 系统总体架构图": Path(r"C:\Users\帆\Desktop\毕业设计\docs\flowcharts\thesis_system_architecture.png"),
    "图3-2 基本蚁群算法流程图": Path(r"C:\Users\帆\Desktop\毕业设计\docs\flowcharts\thesis_basic_aco_flow.png"),
    "图3-3 自适应参数调整策略流程图": Path(r"C:\Users\帆\Desktop\毕业设计\docs\flowcharts\thesis_adaptive_flow.png"),
    "图3-4 遗传算法增强策略流程图": Path(r"C:\Users\帆\Desktop\毕业设计\docs\flowcharts\thesis_ga_flow.png"),
    "图3-5 2-opt局部搜索策略流程图": Path(r"C:\Users\帆\Desktop\毕业设计\docs\flowcharts\thesis_two_opt_flow.png"),
}

FORMULA_OVERRIDES = {
    r"\min L = \sum_{i=1}^{n} d_{\pi_i, \pi_{i+1}} + d_{\pi_n, \pi_1}": [
        "min L = Σ(i=1→n) d(πᵢ, πᵢ₊₁) + d(πₙ, π₁)"
    ],
    r"\sum_{i=1}^{n} x_{ij} = 1, \quad j = 1, 2, ..., n": [
        "Σ(i=1→n) xᵢⱼ = 1，    j = 1, 2, …, n"
    ],
    r"\sum_{j=1}^{n} x_{ij} = 1, \quad i = 1, 2, ..., n": [
        "Σ(j=1→n) xᵢⱼ = 1，    i = 1, 2, …, n"
    ],
    r"x_{ij} \in \{0, 1\}, \quad i, j = 1, 2, ..., n": [
        "xᵢⱼ ∈ {0, 1}，    i, j = 1, 2, …, n"
    ],
    r"\tau_{ij}(t+1) = (1 - \rho) \cdot \tau_{ij}(t)": [
        "τᵢⱼ(t + 1) = (1 − ρ) · τᵢⱼ(t)"
    ],
    r"\tau_{ij}(t+1) = \tau_{ij}(t+1) + \sum_{k=1}^{m} \Delta\tau_{ij}^k": [
        "τᵢⱼ(t + 1) = τᵢⱼ(t + 1) + Σ(k=1→m) Δτᵢⱼᵏ"
    ],
    r"\Delta\tau_{ij}^k = \begin{cases} Q/L_k, & \text{若蚂蚁}k\text{经过路径}(i,j) \\ 0, & \text{否则} \end{cases}": [
        "Δτᵢⱼᵏ = ⎧ Q / Lₖ，若蚂蚁 k 经过路径 (i, j)",
        "        ⎩ 0，      否则",
    ],
    r"\Delta\tau_{ij}^k = \begin{cases} Q/d_{ij}, & \text{若蚂蚁}k\text{经过路径}(i,j) \\ 0, & \text{否则} \end{cases}": [
        "Δτᵢⱼᵏ = ⎧ Q / dᵢⱼ，若蚂蚁 k 经过路径 (i, j)",
        "        ⎩ 0，      否则",
    ],
    r"\Delta\tau_{ij}^k = \begin{cases} Q, & \text{若蚂蚁}k\text{经过路径}(i,j) \\ 0, & \text{否则} \end{cases}": [
        "Δτᵢⱼᵏ = ⎧ Q，若蚂蚁 k 经过路径 (i, j)",
        "        ⎩ 0， 否则",
    ],
    r"P_{ij}^k = \begin{cases} \frac{[\tau_{ij}]^\alpha \cdot [\eta_{ij}]^\beta}{\sum_{s \in \text{allow}_k} [\tau_{is}]^\alpha \cdot [\eta_{is}]^\beta}, & \text{若}j \in \text{allow}_k \\ 0, & \text{否则} \end{cases}": [
        "Pᵏᵢⱼ = ⎧ ([τᵢⱼ]^α · [ηᵢⱼ]^β) / (Σ(s∈allowₖ) [τᵢₛ]^α · [ηᵢₛ]^β)，若 j ∈ allowₖ",
        "       ⎩ 0， 否则",
    ],
    r"p_i = \frac{f_i}{\sum_{j=1}^{N} f_j}": [
        "pᵢ = fᵢ / Σ(j=1→N) fⱼ"
    ],
    r"L' = L - d_{\pi_i, \pi_{i+1}} - d_{\pi_j, \pi_{j+1}} + d_{\pi_i, \pi_j} + d_{\pi_{i+1}, \pi_{j+1}}": [
        "L′ = L − d(πᵢ, πᵢ₊₁) − d(πⱼ, πⱼ₊₁) + d(πᵢ, πⱼ) + d(πᵢ₊₁, πⱼ₊₁)"
    ],
    r"\Delta L = L' - L = d_{\pi_i, \pi_j} + d_{\pi_{i+1}, \pi_{j+1}} - d_{\pi_i, \pi_{i+1}} - d_{\pi_j, \pi_{j+1}}": [
        "ΔL = L′ − L = d(πᵢ, πⱼ) + d(πᵢ₊₁, πⱼ₊₁) − d(πᵢ, πᵢ₊₁) − d(πⱼ, πⱼ₊₁)"
    ],
    r"\alpha(t) = \alpha_{min} + (\alpha_{max} - \alpha_{min}) \times \frac{t}{T}": [
        "α(t) = αₘᵢₙ + (αₘₐₓ − αₘᵢₙ) × t / T"
    ],
    r"\beta(t) = \beta_{max} - (\beta_{max} - \beta_{min}) \times \frac{t}{T}": [
        "β(t) = βₘₐₓ − (βₘₐₓ − βₘᵢₙ) × t / T"
    ],
}


def set_run_font(run, font_name: str, size: Pt | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_borders(paragraph, color: str = "D9D9D9"):
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def clear_default_paragraph(paragraph):
    for run in paragraph.runs:
        run.clear()
    if paragraph.text:
        paragraph.text = ""


def set_paragraph_format(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: Pt | None = Pt(24),
    line_spacing: float = 1.5,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
    left_indent: Pt | None = None,
):
    fmt = paragraph.paragraph_format
    paragraph.alignment = alignment
    fmt.first_line_indent = first_line_indent
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.line_spacing = line_spacing
    fmt.space_before = space_before
    fmt.space_after = space_after
    fmt.left_indent = left_indent
    fmt.right_indent = None


def add_inline_runs(paragraph, text: str, *, font_name: str, font_size: Pt, default_bold: bool = False):
    # Supports **bold** and `inline code`.
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    pos = 0

    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, font_name, font_size, bold=default_bold)

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, font_name, font_size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, CODE_FONT, CODE_SIZE)
        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, font_name, font_size, bold=default_bold)


def is_horizontal_rule(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and set(stripped) <= {"-", "*", "_"} and len(stripped) >= 3


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not is_table_line(stripped):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return all(re.fullmatch(r"[:\- ]+", cell or "-") for cell in cells)


def is_toc_entry(line: str) -> bool:
    stripped = line.strip()
    if re.match(r"^\d+(?:\.\d+)*\s+", stripped):
        return True
    if re.match(r"^第[一二三四五六七八九十]+章", stripped):
        return True
    return False


def is_list_item(line: str) -> bool:
    stripped = line.strip()
    return bool(
        re.match(r"^[-*]\s+", stripped)
        or re.match(r"^\d+\.\s+", stripped)
        or re.match(r"^[（(]\d+[）)]", stripped)
    )


def heading_level(line: str) -> int | None:
    match = re.match(r"^(#{1,4})\s+(.*)$", line)
    if not match:
        return None
    return len(match.group(1))


def strip_list_marker(line: str) -> str:
    stripped = line.strip()
    stripped = re.sub(r"^[-*]\s+", "", stripped)
    stripped = re.sub(r"^\d+\.\s+", "", stripped)
    return stripped


def add_title(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        line_spacing=1.5,
        space_before=Pt(0),
        space_after=Pt(12),
    )
    run = p.add_run(text.strip())
    set_run_font(run, HEADING_FONT, TITLE_SIZE, bold=True)


def add_subtitle(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        line_spacing=1.5,
        space_before=Pt(0),
        space_after=Pt(12),
    )
    run = p.add_run(text.strip())
    set_run_font(run, HEADING_FONT, SUBTITLE_SIZE, bold=True)


def add_heading(document: Document, text: str, level: int):
    p = document.add_paragraph()
    centered_titles = {"摘要", "Abstract", "目录"}
    alignment = WD_ALIGN_PARAGRAPH.CENTER if text.strip() in centered_titles else WD_ALIGN_PARAGRAPH.LEFT
    before = Pt(12) if level <= 2 else Pt(6)
    after = Pt(6)
    set_paragraph_format(
        p,
        alignment=alignment,
        first_line_indent=None,
        line_spacing=1.5,
        space_before=before,
        space_after=after,
    )
    run = p.add_run(text.strip())
    set_run_font(run, HEADING_FONT, HEADING_SIZE, bold=True)


def add_body_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Pt(24),
        line_spacing=1.5,
        space_before=Pt(0),
        space_after=Pt(0),
    )
    add_inline_runs(p, text.strip(), font_name=BODY_FONT, font_size=BODY_SIZE)


def add_toc_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=1.5,
        space_before=Pt(0),
        space_after=Pt(0),
        left_indent=Pt(12),
    )
    add_inline_runs(p, text.strip(), font_name=BODY_FONT, font_size=BODY_SIZE)


def add_list_paragraph(document: Document, text: str, ordered: bool):
    p = document.add_paragraph(style="List Number" if ordered else "List Bullet")
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=None,
        line_spacing=1.5,
        space_before=Pt(0),
        space_after=Pt(0),
    )
    add_inline_runs(p, text.strip(), font_name=BODY_FONT, font_size=BODY_SIZE)


def add_enumerated_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=None,
        line_spacing=1.5,
        space_before=Pt(0),
        space_after=Pt(0),
    )
    add_inline_runs(p, text.strip(), font_name=BODY_FONT, font_size=BODY_SIZE)


def get_formula_lines(content: str) -> list[str]:
    if content in FORMULA_OVERRIDES:
        return FORMULA_OVERRIDES[content]

    fallback = content
    replacements = [
        (r"\alpha", "α"),
        (r"\beta", "β"),
        (r"\tau", "τ"),
        (r"\Delta", "Δ"),
        (r"\rho", "ρ"),
        (r"\eta", "η"),
        (r"\pi", "π"),
        (r"\cdot", "·"),
        (r"\times", "×"),
        (r"\in", "∈"),
        (r"\neq", "≠"),
        (r"\sum", "Σ"),
        (r"\min", "min"),
        (r"\quad", "    "),
        (r"\text{若}", "若"),
        (r"\text{否则}", "否则"),
    ]

    for old, new in replacements:
        fallback = fallback.replace(old, new)

    fallback = fallback.replace("{", "").replace("}", "")
    fallback = fallback.replace("\\", "")
    fallback = re.sub(r"\s+", " ", fallback).strip()
    return [fallback]


def add_formula_paragraph(document: Document, text: str):
    content = text.strip()
    if content.startswith("$$") and content.endswith("$$"):
        content = content[2:-2].strip()

    lines = get_formula_lines(content)
    last_index = len(lines) - 1

    for index, line in enumerate(lines):
        p = document.add_paragraph()
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=None,
            line_spacing=1.5,
            space_before=Pt(6) if index == 0 else Pt(0),
            space_after=Pt(6) if index == last_index else Pt(0),
        )
        run = p.add_run(line)
        set_run_font(run, MATH_FONT, BODY_SIZE)


def add_image_paragraph(document: Document, image_path: Path, width_cm: float = 15.5):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        line_spacing=1.5,
        space_before=Pt(6),
        space_after=Pt(6),
    )
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def add_code_block(document: Document, lines: list[str], language: str | None):
    if language:
        label = document.add_paragraph()
        set_paragraph_format(
            label,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            line_spacing=1.5,
            space_before=Pt(6),
            space_after=Pt(3),
        )
        run = label.add_run(f"{language} 代码")
        set_run_font(run, BODY_FONT, BODY_SIZE, bold=True)

    for raw_line in lines:
        p = document.add_paragraph()
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            line_spacing=1.0,
            space_before=Pt(0),
            space_after=Pt(0),
            left_indent=Pt(18),
        )
        run = p.add_run(raw_line.rstrip("\n").replace("\t", "    "))
        set_run_font(run, CODE_FONT, CODE_SIZE)
        set_paragraph_borders(p)

    spacer = document.add_paragraph()
    set_paragraph_format(
        spacer,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=1.0,
        space_before=Pt(0),
        space_after=Pt(6),
    )


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document: Document, lines: list[str]):
    rows = [split_table_row(line) for line in lines if not is_table_separator(line)]
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row in enumerate(rows):
        for j in range(col_count):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = row[j] if j < len(row) else ""

            if i == 0:
                set_cell_shading(cell, "EDEDED")

            for paragraph in cell.paragraphs:
                clear_default_paragraph(paragraph)
                set_paragraph_format(
                    paragraph,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=None,
                    line_spacing=1.5,
                    space_before=Pt(0),
                    space_after=Pt(0),
                )
                add_inline_runs(
                    paragraph,
                    row[j] if j < len(row) else "",
                    font_name=BODY_FONT,
                    font_size=BODY_SIZE,
                    default_bold=(i == 0),
                )

    document.add_paragraph()


def setup_document(document: Document):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def convert_markdown(input_path: Path, output_path: Path):
    lines = input_path.read_text(encoding="utf-8").splitlines()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_document(doc)

    in_code_block = False
    code_lines: list[str] = []
    code_language: str | None = None
    table_lines: list[str] = []
    title_done = False
    subtitle_done = False

    i = 0
    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("**") and stripped.endswith("**"):
            bold_text = stripped.strip("*").strip()
            if bold_text in FLOWCHART_IMAGE_MAP and FLOWCHART_IMAGE_MAP[bold_text].exists():
                # Skip the following ASCII art fenced block if present.
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and lines[j].strip().startswith("```"):
                    j += 1
                    while j < len(lines) and not lines[j].strip().startswith("```"):
                        j += 1
                    if j < len(lines):
                        j += 1
                add_image_paragraph(doc, FLOWCHART_IMAGE_MAP[bold_text])
                i = j
                continue

        if in_code_block:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines, code_language)
                code_lines = []
                code_language = None
                in_code_block = False
            else:
                code_lines.append(raw_line)
            i += 1
            continue

        if stripped.startswith("```"):
            in_code_block = True
            code_language = stripped[3:].strip() or None
            i += 1
            continue

        if is_table_line(stripped):
            table_lines = [stripped]
            i += 1
            while i < len(lines) and is_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines)
            continue

        if not stripped:
            i += 1
            continue

        if is_horizontal_rule(stripped):
            i += 1
            continue

        if stripped.startswith("$$") and stripped.endswith("$$"):
            add_formula_paragraph(doc, stripped)
            i += 1
            continue

        level = heading_level(stripped)
        if level is not None:
            text = re.sub(r"^#{1,4}\s+", "", stripped).strip()
            if level == 1 and not title_done:
                add_title(doc, text)
                title_done = True
            elif level == 2 and text == "毕业设计说明书" and not subtitle_done:
                add_subtitle(doc, text)
                subtitle_done = True
            else:
                add_heading(doc, text, level)
            i += 1
            continue

        if is_toc_entry(stripped):
            add_toc_paragraph(doc, stripped)
            i += 1
            continue

        if re.match(r"^[（(]\d+[）)]", stripped):
            add_enumerated_paragraph(doc, stripped)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_list_paragraph(doc, strip_list_marker(stripped), ordered=True)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            add_list_paragraph(doc, strip_list_marker(stripped), ordered=False)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            set_paragraph_format(
                p,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=None,
                line_spacing=1.5,
                space_before=Pt(0),
                space_after=Pt(0),
            )
            add_inline_runs(p, stripped, font_name=BODY_FONT, font_size=BODY_SIZE)
            i += 1
            continue

        add_body_paragraph(doc, stripped)
        i += 1

    if in_code_block and code_lines:
        add_code_block(doc, code_lines, code_language)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Convert thesis markdown to formatted Word document.")
    parser.add_argument("input", type=Path, help="Input markdown file")
    parser.add_argument("output", type=Path, help="Output docx file")
    args = parser.parse_args()

    convert_markdown(args.input, args.output)


if __name__ == "__main__":
    main()
